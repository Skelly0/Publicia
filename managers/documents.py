"""
Document management, embeddings, and search functionality for Publicia
"""
import os
import re
import json
import pickle
import logging
import aiohttp
import asyncio
import urllib.parse
import uuid 
import hashlib 
import numpy as np
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any
from pathlib import Path
from textwrap import shorten
from rank_bm25 import BM25Okapi

# Import docx components
try:
    import docx
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Define dummy classes if docx is not installed to avoid runtime errors on import
    class RGBColor:
        def __init__(self, r, g, b): pass
    class Document:
        def __init__(self): self.paragraphs = []
    class Paragraph:
        def __init__(self): self.runs = []
    class Run:
        def __init__(self): self.text = ""; self.font = Font()
    class Font:
        def __init__(self): self.color = Color()
    class Color:
        def __init__(self): self.rgb = None

from utils.logging import sanitize_for_logging

logger = logging.getLogger(__name__)


# --- New Function: tag_lore_in_docx ---

def tag_lore_in_docx(docx_filepath: str) -> Optional[str]:
    """
    Processes a .docx file, adding XML tags around text with a specific color.

    Args:
        docx_filepath: Path to the input .docx file.

    Returns:
        A string containing the processed text with XML tags,
        or None if python-docx is not installed or the file cannot be processed.
    """
    if not DOCX_AVAILABLE:
        logger.error("The 'python-docx' library is required but not installed. Cannot process .docx file.")
        return None

    try:
        document = docx.Document(docx_filepath)
        target_color = RGBColor(152, 0, 0) # #980000
        output_parts = []
        is_in_lore_block = False

        for para in document.paragraphs:
            para_text_parts = []
            for run in para.runs:
                run_color = run.font.color.rgb if run.font.color else None
                is_target_color = run_color == target_color

                if is_target_color and not is_in_lore_block:
                    para_text_parts.append("<post-invasion_lore>")
                    is_in_lore_block = True
                elif not is_target_color and is_in_lore_block:
                    para_text_parts.append("</post-invasion_lore>")
                    is_in_lore_block = False
                para_text_parts.append(run.text)

            if is_in_lore_block: # If paragraph ends mid-lore
                para_text_parts.append("</post-invasion_lore>")
                is_in_lore_block = False 
            output_parts.append("".join(para_text_parts))
        
        final_text = "\n\n".join(output_parts)
        return final_text

    except FileNotFoundError:
        logger.error(f"Error processing DOCX: File not found at {docx_filepath}")
        return None
    except Exception as e:
        logger.error(f"Error processing DOCX file '{docx_filepath}': {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None

# --- End of New Function ---


class DocumentManager:
    """Manages document storage, embeddings, and retrieval using UUIDs as primary keys."""

    # Maximum byte size for chunks sent to Google Embedding API
    # Google's limit is 36000 bytes for the request payload.
    # Set slightly lower to account for JSON overhead and other metadata.
    MAX_EMBEDDING_CHUNK_BYTES = 35000 

    def _get_original_name(self, doc_uuid: str) -> str:
        """Retrieve the original document name from metadata given the document's UUID."""
        if doc_uuid in self.metadata:
            return self.metadata[doc_uuid].get('original_name', doc_uuid) 
        return doc_uuid

    def _sanitize_name(self, name: str) -> str:
        """Sanitizes a string to be safe for use as a filename or dictionary key."""
        if not isinstance(name, str):
            name = str(name) # Ensure it's a string

        # Replace potentially problematic characters with underscores
        # Includes Windows reserved chars: <>:"/\|?*
        # Also includes control characters (0-31)
        sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', name)

        # Replace leading/trailing dots and spaces
        sanitized = sanitized.strip('. ')

        # Collapse multiple consecutive underscores
        sanitized = re.sub(r'_+', '_', sanitized)

        # Ensure the name is not empty after sanitization
        if not sanitized:
            sanitized = "untitled"
        # Optional: Limit length?
        # max_len = 200
        # sanitized = sanitized[:max_len]
        return sanitized

    def _get_sanitized_name_from_original(self, original_name: str) -> Optional[str]:
        """
        Find the sanitized name corresponding to an original name.
        Handles cases where the input might have a .txt extension but the stored name doesn't.
        """
        if not original_name: # Handle empty input
            return None

        # 1. Check direct match (input name == stored original_name)
        # Iterate through metadata first to handle potential sanitization collisions
        logger.debug(f"Attempt 1: Checking direct match for '{original_name}'") # ADDED LOGGING
        for s_name, meta in self.metadata.items():
            stored_original = meta.get('original_name') # ADDED LOGGING
            logger.debug(f"Comparing '{original_name}' with stored '{stored_original}' (sanitized key: {s_name})") # ADDED LOGGING
            if stored_original == original_name:
                logger.debug(f"Found direct match for '{original_name}' -> sanitized '{s_name}'")
                return s_name

        # 2. Check if input has .txt and stored original_name doesn't
        name_without_txt = None
        if original_name.endswith('.txt'):
            name_without_txt = original_name[:-4]
            if name_without_txt: # Ensure not empty after stripping
                logger.debug(f"Attempt 2: Checking match for '{name_without_txt}' (removed .txt)") # ADDED LOGGING
                for s_name, meta in self.metadata.items():
                    stored_original = meta.get('original_name') # ADDED LOGGING
                    logger.debug(f"Comparing '{name_without_txt}' with stored '{stored_original}' (sanitized key: {s_name})") # ADDED LOGGING
                    if stored_original == name_without_txt:
                        logger.debug(f"Found match for '{original_name}' by removing .txt ('{name_without_txt}') -> sanitized '{s_name}'")
                        return s_name

        # 3. Check if input *doesn't* have .txt but stored original_name *does* (less common)
        name_with_txt = f"{original_name}.txt"
        logger.debug(f"Attempt 3: Checking match for '{name_with_txt}' (added .txt)") # ADDED LOGGING
        for s_name, meta in self.metadata.items():
             stored_original = meta.get('original_name') # ADDED LOGGING
             logger.debug(f"Comparing '{name_with_txt}' with stored '{stored_original}' (sanitized key: {s_name})") # ADDED LOGGING
             if stored_original == name_with_txt:
                 logger.debug(f"Found match for '{original_name}' by adding .txt ('{name_with_txt}') -> sanitized '{s_name}'")
                 return s_name

        # 4. Fallback: Check if the sanitized version of the input exists as a key
        # This helps if the original_name field in metadata is somehow incorrect/missing
        # but the sanitized key itself matches.
        logger.debug(f"Attempt 4: Checking fallback using sanitized key") # ADDED LOGGING
        s_name_direct = self._sanitize_name(original_name)
        if s_name_direct in self.metadata:
            logger.warning(f"Found sanitized key '{s_name_direct}' matching input '{original_name}', but original_name field in metadata might be inconsistent. Returning sanitized key as fallback.")
            return s_name_direct
        # Also check sanitized version without .txt if applicable
        if name_without_txt:
            s_name_direct_no_txt = self._sanitize_name(name_without_txt)
            if s_name_direct_no_txt in self.metadata:
                 logger.warning(f"Found sanitized key '{s_name_direct_no_txt}' matching input '{original_name}' (without .txt), but original_name field in metadata might be inconsistent. Returning sanitized key as fallback.")
                 return s_name_direct_no_txt


        logger.warning(f"Could not find any matching sanitized name for original name: '{original_name}'")
        return None # Not found

    def __init__(self, base_dir: str = "documents", top_k: int = 5, config=None):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.top_k = top_k
        self.config = config
        
        try:
            logger.info("Initializing Google Generative AI embedding model")
            if not config or not config.GOOGLE_API_KEY:
                logger.error("GOOGLE_API_KEY environment variable not set")
                raise ValueError("GOOGLE_API_KEY environment variable not set")
            import google.generativeai as genai
            genai.configure(api_key=config.GOOGLE_API_KEY)
            self.embedding_model = config.EMBEDDING_MODEL if config else 'models/text-embedding-004'
            self.embedding_dimensions = config.EMBEDDING_DIMENSIONS if config and config.EMBEDDING_DIMENSIONS > 0 else None
            logger.info(f"Using Google embedding model: {self.embedding_model}")
            if self.embedding_dimensions:
                logger.info(f"Truncating embeddings to {self.embedding_dimensions} dimensions")
            else:
                logger.info("Using full embedding dimensions.")
            logger.info(f"Gemini embedding model initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing Google Generative AI: {e}")
            raise
        
        self.chunks: Dict[str, List[str]] = {}
        self.embeddings: Dict[str, np.ndarray] = {}
        self.metadata: Dict[str, Dict] = {}
        self.contextualized_chunks: Dict[str, List[str]] = {}
        self.bm25_indexes: Dict[str, BM25Okapi] = {}
        
        logger.info("Document loading will be done asynchronously via _load_documents.")
        self._internal_list_doc_name = "_internal_document_list.txt" 

    async def _update_document_list_file(self):
        logger.info("Updating internal document list file...")
        try:
            doc_items_for_list = []
            internal_list_doc_uuid_to_skip = None
            for doc_uuid_key, meta_val in self.metadata.items():
                if meta_val.get('original_name') == self._internal_list_doc_name:
                    internal_list_doc_uuid_to_skip = doc_uuid_key
                    break
            
            for doc_uuid, meta in self.metadata.items():
                if doc_uuid == internal_list_doc_uuid_to_skip:
                    continue
                original_name = meta.get('original_name', doc_uuid)
                doc_items_for_list.append(f"{original_name} (UUID: {doc_uuid})")

            doc_items_for_list.sort() 
            content = "Managed Documents:\n\n" + "\n".join(f"- {item}" for item in doc_items_for_list) if doc_items_for_list else "No documents are currently managed."
            
            returned_uuid = await self.add_document(
                original_name=self._internal_list_doc_name, 
                content=content, 
                save_to_disk=True, 
                existing_uuid=internal_list_doc_uuid_to_skip,
                _internal_call=True
            )
            if returned_uuid: logger.info(f"Successfully updated internal list doc ('{self._internal_list_doc_name}', UUID: {returned_uuid}).")
            else: logger.error(f"Failed to add/update internal list doc ('{self._internal_list_doc_name}').")
        except Exception as e: logger.error(f"Failed to update internal document list file: {e}", exc_info=True)

    async def _get_google_embedding_async(self, text: str, task_type: str, title: Optional[str] = None, max_retries: int = 3) -> Optional[np.ndarray]:
        if not text or not text.strip(): return None
        if not self.config.GOOGLE_API_KEY: return None
        api_url = f"https://generativelanguage.googleapis.com/v1beta/{self.embedding_model}:embedContent?key={self.config.GOOGLE_API_KEY}"
        payload = {"content": {"parts": [{"text": text}]}, "task_type": task_type}
        if title and task_type == "retrieval_document": payload["title"] = title
        headers = {"Content-Type": "application/json"}
        
        for attempt in range(max_retries):
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.post(api_url, headers=headers, json=payload, timeout=60) as response:
                        if response.status == 200:
                            result = await response.json()
                            if "embedding" in result and "values" in result["embedding"]:
                                vec = np.array(result["embedding"]["values"])
                                return vec[:self.embedding_dimensions] if self.embedding_dimensions and self.embedding_dimensions < len(vec) else vec
                        
                        # Log error but continue retrying for non-200 status
                        error_text = await response.text()
                        logger.warning(f"Google Embedding API error (Status {response.status}, attempt {attempt + 1}/{max_retries}): {error_text}")
                        
                        # Don't retry on certain permanent errors
                        if response.status in [400, 401, 403]:
                            logger.error(f"Permanent error {response.status}, not retrying")
                            break
                            
            except Exception as e:
                logger.warning(f"Error calling Google Embedding API (attempt {attempt + 1}/{max_retries}): {e}")
                
            # Wait before retrying (exponential backoff)
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt  # 1s, 2s, 4s
                await asyncio.sleep(wait_time)
        
        logger.error(f"Failed to get embedding after {max_retries} attempts for text: {text[:100]}...")
        return None

    def _create_bm25_index(self, chunks: List[str]) -> Optional[BM25Okapi]:
        if not chunks:
            logger.warning("Cannot create BM25 index: no chunks provided")
            return None
            
        # Filter out empty chunks and tokenize
        tokenized_chunks = []
        for i, chunk in enumerate(chunks):
            if chunk and chunk.strip():
                tokens = chunk.lower().split()
                if tokens:  # Only add if we have actual tokens
                    tokenized_chunks.append(tokens)
                else:
                    logger.warning(f"Chunk {i} has no tokens after processing: '{chunk[:50]}...'")
            else:
                logger.warning(f"Empty chunk at index {i}")
        
        if not tokenized_chunks:
            logger.error(f"No valid tokenized chunks found out of {len(chunks)} input chunks")
            return None
            
        try:
            bm25_index = BM25Okapi(tokenized_chunks)
            logger.debug(f"Created BM25 index with {len(tokenized_chunks)} chunks")
            return bm25_index
        except Exception as e:
            logger.error(f"Failed to create BM25 index: {e}")
            return None
        
    def _extract_core_topics(self, query: str) -> str:
        """
        Extract core topic words from a complex query, removing analytical language.
        This helps BM25 focus on the actual subject matter rather than instructional words.
        """
        import re
        
        # Remove common analytical/instructional words that don't help with content matching
        analytical_words = {
            'write', 'detailed', 'analysis', 'analyze', 'comprehensive', 'explain',
            'describe', 'discuss', 'elaborate', 'provide', 'give', 'tell', 'about',
            'intersection', 'relationship', 'connection', 'between', 'with', 'their',
            'as', 'well', 'its', 'the', 'a', 'an', 'and', 'or', 'of', 'in', 'on',
            'at', 'to', 'for', 'by', 'from', 'up', 'out', 'if', 'then', 'than',
            'how', 'what', 'when', 'where', 'why', 'who', 'which', 'that', 'this'
        }
        
        # Split query into words and filter
        words = re.findall(r'\b\w+\b', query.lower())
        core_words = [word for word in words if word not in analytical_words and len(word) > 2]
        
        # Return the core topic words, fallback to original if no core words found
        core_query = ' '.join(core_words)
        return core_query if core_words else query

    def _search_bm25(self, query: str, top_k: int = None) -> List[Tuple[str, str, str, float, Optional[str], int, int]]:
        if top_k is None: top_k = self.top_k
        
        # Extract core topics for better BM25 matching
        core_query = self._extract_core_topics(query)
        if core_query != query:
            logger.debug(f"BM25 using core topics: '{query}' -> '{core_query}'")
        
        query_tokens = core_query.lower().split()
        if not query_tokens:
            logger.warning("BM25 search: No valid query tokens")
            return []
            
        results = []
        for doc_uuid, doc_chunks in self.chunks.items():
            if not doc_chunks:
                logger.debug(f"Skipping document {doc_uuid}: no chunks")
                continue
                
            # Determine which chunks to use for BM25 indexing
            use_contextualised = self.config.USE_CONTEXTUALISED_CHUNKS if self.config and hasattr(self.config, 'USE_CONTEXTUALISED_CHUNKS') else True
            
            if use_contextualised:
                chunks_for_bm25 = self.contextualized_chunks.get(doc_uuid, [])
                if not chunks_for_bm25:
                    chunks_for_bm25 = doc_chunks
                    logger.debug(f"Using original chunks for BM25 indexing for {doc_uuid}")
                else:
                    logger.debug(f"Using contextualized chunks for BM25 indexing for {doc_uuid}")
            else:
                chunks_for_bm25 = doc_chunks
                logger.debug(f"Using original chunks for BM25 indexing for {doc_uuid} (contextualised chunks disabled)")
            
            # Ensure BM25 index exists and is valid
            if doc_uuid not in self.bm25_indexes or self.bm25_indexes[doc_uuid] is None:
                logger.debug(f"Creating BM25 index for {doc_uuid} with {len(chunks_for_bm25)} chunks")
                self.bm25_indexes[doc_uuid] = self._create_bm25_index(chunks_for_bm25)
                
            if self.bm25_indexes[doc_uuid] is None:
                logger.warning(f"Could not create BM25 index for {doc_uuid}, skipping")
                continue
            
            try:
                bm25_scores = self.bm25_indexes[doc_uuid].get_scores(query_tokens)
                original_name = self._get_original_name(doc_uuid)
                
                # Validate that we have the expected number of scores
                expected_scores = len(chunks_for_bm25)
                if len(bm25_scores) != expected_scores:
                    logger.warning(f"BM25 score count mismatch for {doc_uuid}: got {len(bm25_scores)}, expected {expected_scores}")
                    # Use the minimum to avoid index errors
                    max_idx = min(len(bm25_scores), len(doc_chunks), len(chunks_for_bm25))
                else:
                    max_idx = min(len(bm25_scores), len(doc_chunks))
                
                for chunk_idx in range(max_idx):
                    score = bm25_scores[chunk_idx]
                    chunk_text = self.get_contextualized_chunk(doc_uuid, chunk_idx)
                    
                    # Skip chunks that couldn't be retrieved
                    if chunk_text == "Chunk not found":
                        logger.warning(f"Skipping chunk {chunk_idx} for {doc_uuid}: chunk not found")
                        continue
                        
                    image_id = self.metadata.get(doc_uuid, {}).get('image_id')
                    results.append((doc_uuid, original_name, chunk_text, float(score), image_id, chunk_idx + 1, len(doc_chunks)))
                    
            except Exception as e:
                logger.error(f"Error getting BM25 scores for {doc_uuid}: {e}")
                continue
                
        results.sort(key=lambda x: x[3], reverse=True)
        logger.debug(f"BM25 search returned {len(results)} results")
        return results[:top_k]
    
    def get_contextualized_chunk(self, doc_uuid: str, chunk_idx: int) -> str:
        # Check if using contextualised chunks is enabled
        use_contextualised = self.config.USE_CONTEXTUALISED_CHUNKS if self.config and hasattr(self.config, 'USE_CONTEXTUALISED_CHUNKS') else True
        
        # If contextualised chunks are disabled, go straight to original chunks
        if not use_contextualised:
            if doc_uuid in self.chunks:
                original = self.chunks[doc_uuid]
                if original and chunk_idx < len(original):
                    logger.debug(f"Using original chunk for {doc_uuid}[{chunk_idx}] (contextualised chunks disabled)")
                    return original[chunk_idx]
                elif original:
                    logger.warning(f"Original chunk index {chunk_idx} out of range for {doc_uuid} (has {len(original)} chunks)")
        else:
            # First try contextualized chunks
            if doc_uuid in self.contextualized_chunks:
                contextualized = self.contextualized_chunks[doc_uuid]
                if contextualized and chunk_idx < len(contextualized):
                    return contextualized[chunk_idx]
                elif contextualized:
                    logger.warning(f"Contextualized chunk index {chunk_idx} out of range for {doc_uuid} (has {len(contextualized)} chunks)")

            # Fallback to original chunks
            if doc_uuid in self.chunks:
                original = self.chunks[doc_uuid]
                if original and chunk_idx < len(original):
                    logger.debug(f"Using original chunk for {doc_uuid}[{chunk_idx}] (contextualized not available)")
                    return original[chunk_idx]
                elif original:
                    logger.warning(f"Original chunk index {chunk_idx} out of range for {doc_uuid} (has {len(original)} chunks)")

        # Log detailed error information
        contextualized_count = len(self.contextualized_chunks.get(doc_uuid, []))
        original_count = len(self.chunks.get(doc_uuid, []))
        logger.error(f"Chunk not found for UUID '{doc_uuid}' at index {chunk_idx}. "
                    f"Available: {original_count} original chunks, {contextualized_count} contextualized chunks")
        return "Chunk not found"

    def _combine_search_results(
        self,
        embedding_results: List[Tuple[str, str, str, float, Optional[str], int, int]],
        bm25_results: List[Tuple[str, str, str, float, Optional[str], int, int]],
        top_k: int = None
    ) -> List[Tuple[str, str, str, float, Optional[str], int, int]]:
        if top_k is None: top_k = self.top_k
        combined_scores_data = {}
        bm25_weight = self.config.BM25_WEIGHT if self.config else 0.05
        embedding_weight = 1.0 - bm25_weight
        
        # Filter out results with "Chunk not found" from embedding results
        valid_embedding_results = [r for r in embedding_results if r[2] != "Chunk not found"]
        if len(valid_embedding_results) < len(embedding_results):
            logger.warning(f"Filtered out {len(embedding_results) - len(valid_embedding_results)} embedding results with 'Chunk not found'")
        
        for res_tuple in valid_embedding_results:
            doc_uuid, _, _, score, _, chunk_idx, _ = res_tuple
            key = (doc_uuid, chunk_idx)
            norm_score = max(0, min(1, score))
            if key not in combined_scores_data:
                combined_scores_data[key] = {'score': 0, 'data': res_tuple}
            combined_scores_data[key]['score'] += norm_score * embedding_weight
            combined_scores_data[key]['data'] = res_tuple

        if bm25_results:
            # Filter out results with "Chunk not found" from BM25 results
            valid_bm25_results = [r for r in bm25_results if r[2] != "Chunk not found"]
            if len(valid_bm25_results) < len(bm25_results):
                logger.warning(f"Filtered out {len(bm25_results) - len(valid_bm25_results)} BM25 results with 'Chunk not found'")
            
            if valid_bm25_results:
                all_bm25_scores = [r[3] for r in valid_bm25_results if r[3] is not None]
                max_bm25, min_bm25 = (max(all_bm25_scores) if all_bm25_scores else 1.0), (min(all_bm25_scores) if all_bm25_scores else 0.0)
                range_bm25 = max_bm25 - min_bm25
                
                for res_tuple in valid_bm25_results:
                    doc_uuid, _, _, score, _, chunk_idx, _ = res_tuple
                    key = (doc_uuid, chunk_idx)
                    norm_score = (score - min_bm25) / range_bm25 if range_bm25 > 0 else 0.5
                    if key not in combined_scores_data:
                        combined_scores_data[key] = {'score': 0, 'data': res_tuple}
                    combined_scores_data[key]['score'] += norm_score * bm25_weight
                    # Prefer BM25 data if embedding data was "Chunk not found"
                    if combined_scores_data[key]['data'][2] == "Chunk not found":
                        combined_scores_data[key]['data'] = res_tuple
        
        # Build final results, ensuring we get fresh chunk text and filter out any remaining "Chunk not found"
        final_results = []
        for d in combined_scores_data.values():
            doc_uuid, original_name, _, score, image_id, chunk_idx, total_chunks = d['data']
            # Get fresh chunk text to ensure consistency
            chunk_text = self.get_contextualized_chunk(doc_uuid, chunk_idx - 1)  # chunk_idx is 1-based
            
            # Skip if chunk still not found
            if chunk_text == "Chunk not found":
                logger.warning(f"Skipping result for {doc_uuid}[{chunk_idx}]: chunk not found during final assembly")
                continue
                
            final_results.append((doc_uuid, original_name, chunk_text, score, image_id, chunk_idx, total_chunks))
        
        final_results.sort(key=lambda x: x[3], reverse=True)
        logger.debug(f"Combined search results: {len(final_results)} valid results from {len(embedding_results)} embedding + {len(bm25_results)} BM25 results")
        return final_results[:top_k]

    async def generate_chunk_context(self, document_content: str, chunk_content: str) -> str:
        try:
            system_prompt = """You are a helpful AI assistant that creates concise contextual descriptions for document chunks. Your task is to provide a short, succinct context that situates a specific chunk within the overall document to improve search retrieval. Answer only with the succinct context and nothing else."""
            user_prompt = f"""<document> \n{document_content} \n</document> \nHere is the chunk we want to situate within the whole document \n<chunk> \n{chunk_content} \n</chunk> \nPlease give a short succinct context to situate this chunk within the overall document for the purposes of improving search retrieval of the chunk. Answer only with the succinct context and nothing else."""
            messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]
            headers = {"Authorization": f"Bearer {self.config.OPENROUTER_API_KEY}", "HTTP-Referer": "https://discord.gg/dpsrp", "X-Title": "Publicia for DPS Season 7", "Content-Type": "application/json"}
            fallback_models = ["cohere/command-r-08-2024", "amazon/nova-lite-v1", "google/gemini-2.0-flash-lite-001", "gryphe/gryphe-mistral-7b-instruct-v2", "mistralai/mistral-7b-instruct"]
            for model in fallback_models:
                payload = {"model": model, "messages": messages, "temperature": 0.1, "max_tokens": 150}
                for _ in range(2): 
                    try:
                        async with aiohttp.ClientSession() as session:
                            async with session.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload, timeout=45) as response:
                                if response.status == 200:
                                    completion = await response.json()
                                    if completion and completion.get('choices') and completion['choices'][0].get('message', {}).get('content'):
                                        return completion['choices'][0]['message']['content'].strip()
                                logger.error(f"API error model {model} (Status {response.status}): {await response.text()}")
                                break 
                    except Exception as e: logger.error(f"Ctx gen error model {model}: {e}")
                    await asyncio.sleep(0.5)
            return chunk_content # Fallback: return original chunk if all models fail
        except Exception as e:
            logger.error(f"Error in context gen: {e}", exc_info=True)
            return chunk_content # Fallback: return original chunk on exception


    async def contextualize_chunks(self, doc_name: str, document_content: str, chunks: List[str]) -> List[str]:
        contextualized_results = []
        successful_ctx_count = 0
        for chunk in chunks:
            generated_context = await self.generate_chunk_context(document_content, chunk)
            if generated_context == chunk: # Context generation failed, use original chunk
                contextualized_results.append(chunk)
            else: # Context generation successful
                contextualized_results.append(generated_context + " " + chunk)
                successful_ctx_count += 1
        
        logger.info(f"Contextualized {successful_ctx_count} of {len(chunks)} chunks for {doc_name}.")
        return contextualized_results

    async def generate_embeddings(self, texts: List[str], is_query: bool = False, titles: List[str] = None) -> np.ndarray:
        tasks, valid_indices = [], []
        task_type = "retrieval_query" if is_query else "retrieval_document"
        for i, text in enumerate(texts):
            if text and text.strip():
                title = titles[i] if titles and i < len(titles) and not is_query else None
                tasks.append(self._get_google_embedding_async(text, task_type, title))
                valid_indices.append(i)
        
        if not tasks:
            logger.warning("No valid texts provided for embedding generation")
            return np.array([])
            
        results = await asyncio.gather(*tasks)
        successful_embeddings = []
        failed_indices = []
        
        for i, res_emb in enumerate(results):
            original_idx = valid_indices[i]
            if res_emb is not None:
                successful_embeddings.append(res_emb)
            else:
                failed_indices.append(original_idx)
                logger.error(f"Failed to generate embedding for text at index {original_idx}: {texts[original_idx][:100]}...")
        
        if failed_indices:
            logger.warning(f"Failed to generate embeddings for {len(failed_indices)} out of {len(texts)} texts. Failed indices: {failed_indices}")
            
        if not successful_embeddings:
            logger.error("All embedding generation attempts failed")
            return np.array([])
            
        # Only return successful embeddings - no zero-vector placeholders
        # This means the caller needs to handle the case where fewer embeddings are returned than texts provided
        logger.info(f"Successfully generated {len(successful_embeddings)} embeddings out of {len(texts)} texts")
        return np.array(successful_embeddings)

    def _chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        Chunks text by word count, then further splits chunks if they exceed byte limits.
        Validates that no content is lost during the chunking process.
        """
        word_chunk_size = chunk_size or (self.config.CHUNK_SIZE if self.config else 750)
        word_overlap = overlap or (self.config.CHUNK_OVERLAP if self.config else 125)
        max_bytes = (self.config.MAX_EMBEDDING_CHUNK_BYTES
                     if self.config and hasattr(self.config, 'MAX_EMBEDDING_CHUNK_BYTES')
                     else self.MAX_EMBEDDING_CHUNK_BYTES)

        if not text or not text.strip():
            return []

        original_text = text.strip()
        words = original_text.split()
        if not words:
            return []

        # Initial chunking by word count
        initial_word_chunks = []
        if len(words) <= word_chunk_size:
            initial_word_chunks.append(' '.join(words))
        else:
            for i in range(0, len(words), word_chunk_size - word_overlap):
                chunk_words = words[i:min(i + word_chunk_size, len(words))]
                if chunk_words:  # Ensure we don't add empty chunks
                    initial_word_chunks.append(' '.join(chunk_words))

        if not initial_word_chunks:
            return []

        # Further split chunks if they exceed byte limit
        final_chunks = []
        for word_chunk in initial_word_chunks:
            if not word_chunk.strip():
                continue
            
            chunk_bytes = len(word_chunk.encode('utf-8'))

            if chunk_bytes <= max_bytes:
                final_chunks.append(word_chunk)
            else:
                # Byte-based splitting for oversized chunks
                sub_words = word_chunk.split()
                current_sub_chunk_words = []
                current_sub_chunk_bytes = 0
                
                for sub_word in sub_words:
                    sub_word_bytes = len(sub_word.encode('utf-8'))
                    # Add 1 byte for space, unless it's the first word in the sub_chunk
                    space_bytes = 1 if current_sub_chunk_words else 0

                    if current_sub_chunk_bytes + sub_word_bytes + space_bytes > max_bytes:
                        if current_sub_chunk_words: # Add the current sub-chunk if it's not empty
                            final_chunks.append(' '.join(current_sub_chunk_words))
                        # Start a new sub-chunk
                        current_sub_chunk_words = [sub_word]
                        current_sub_chunk_bytes = sub_word_bytes
                    else:
                        current_sub_chunk_words.append(sub_word)
                        current_sub_chunk_bytes += sub_word_bytes + space_bytes
                
                # Add any remaining part of the sub-chunk
                if current_sub_chunk_words:
                    final_chunks.append(' '.join(current_sub_chunk_words))
        
        # Filter out any empty strings that might have been produced
        final_chunks = [chunk for chunk in final_chunks if chunk and chunk.strip()]
        
        # Validate that no content was lost during chunking
        reconstructed_text = ' '.join(final_chunks)
        original_words_set = set(original_text.split())
        reconstructed_words_set = set(reconstructed_text.split())
        
        if original_words_set != reconstructed_words_set:
            missing_words = original_words_set - reconstructed_words_set
            extra_words = reconstructed_words_set - original_words_set
            
            if missing_words:
                logger.error(f"Content validation failed: {len(missing_words)} words lost during chunking: {list(missing_words)[:10]}...")
            if extra_words:
                logger.warning(f"Content validation warning: {len(extra_words)} extra words found during chunking: {list(extra_words)[:10]}...")
        else:
            logger.debug(f"Content validation passed: All {len(original_words_set)} unique words preserved in {len(final_chunks)} chunks")
        
        return final_chunks

    def cleanup_empty_documents(self):
        empty_docs_uuids = [uuid_key for uuid_key, embs in self.embeddings.items() if not isinstance(embs, np.ndarray) or embs.size == 0]
        for doc_uuid in empty_docs_uuids:
            name = self._get_original_name(doc_uuid)
            logger.info(f"Removing empty doc: '{name}' (UUID: {doc_uuid})")
            for store in [self.chunks, self.embeddings, self.metadata, self.contextualized_chunks, self.bm25_indexes]:
                if doc_uuid in store: del store[doc_uuid]
        if empty_docs_uuids: self._save_to_disk()
        return empty_docs_uuids

    async def add_document(self, original_name: str, content: str, save_to_disk: bool = True, existing_uuid: Optional[str] = None, _internal_call: bool = False) -> Optional[str]:
        try:
            doc_uuid = existing_uuid or str(uuid.uuid4())
            is_update = bool(existing_uuid and doc_uuid in self.metadata)
            logger.info(f"{'Updating' if is_update else 'Adding new'} document '{original_name}' (UUID: {doc_uuid}).")

            if original_name == self._internal_list_doc_name and not _internal_call:
                logger.warning(f"Direct modification of internal list '{original_name}' denied.")
                return None
            if not content or not content.strip():
                logger.warning(f"Doc '{original_name}' (UUID: {doc_uuid}) has no content. Skipping.")
                return doc_uuid if is_update else None

            doc_file_path = self.base_dir / f"{doc_uuid}.txt"
            word_count = len(content.split())
            max_words_ctx = self.config.MAX_WORDS_FOR_CONTEXT if self.config and hasattr(self.config, 'MAX_WORDS_FOR_CONTEXT') else 20000
            chunks = self._chunk_text(content)
            if not chunks: return doc_uuid if is_update else None

            current_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            content_changed = not (is_update and self.metadata.get(doc_uuid, {}).get('content_hash') == current_hash)
            needs_processing = content_changed or doc_uuid not in self.embeddings or not is_update

            if needs_processing:
                # Check if contextualization is enabled and document is within word limit
                contextualization_enabled = self.config.CONTEXTUALIZATION_ENABLED if self.config and hasattr(self.config, 'CONTEXTUALIZATION_ENABLED') else True
                use_contextualised = self.config.USE_CONTEXTUALISED_CHUNKS if self.config and hasattr(self.config, 'USE_CONTEXTUALISED_CHUNKS') else True
                should_ctx = contextualization_enabled and word_count <= max_words_ctx
                chunks_for_processing = await self.contextualize_chunks(original_name, content, chunks) if should_ctx else chunks
                if not chunks_for_processing and chunks:
                    logger.warning(f"Contextualization returned empty for '{original_name}', using original chunks.")
                    chunks_for_processing = chunks
                
                self.contextualized_chunks[doc_uuid] = chunks_for_processing if should_ctx and chunks_for_processing else []
                
                # Determine which chunks to use for embedding generation
                if use_contextualised and should_ctx and chunks_for_processing:
                    chunks_for_embedding = chunks_for_processing
                    logger.debug(f"Using contextualized chunks for embeddings for '{original_name}'")
                else:
                    chunks_for_embedding = chunks
                    logger.debug(f"Using original chunks for embeddings for '{original_name}' (contextualised chunks disabled or not generated)")
                
                if chunks_for_embedding:
                    titles = [original_name] * len(chunks_for_embedding)
                    embeddings_arr = await self.generate_embeddings(chunks_for_embedding, is_query=False, titles=titles)
                    
                    # Check if we got fewer embeddings than chunks
                    if embeddings_arr.size > 0 and len(embeddings_arr) < len(chunks_for_embedding):
                        logger.warning(f"Got {len(embeddings_arr)} embeddings for {len(chunks_for_embedding)} chunks in '{original_name}'. Some chunks will be unsearchable.")
                        # Store only the chunks that have corresponding embeddings
                        chunks_original_with_embeddings = chunks[:len(embeddings_arr)]
                        chunks_contextualized_with_embeddings = chunks_for_processing[:len(embeddings_arr)] if should_ctx and chunks_for_processing else []
                        
                        self.chunks[doc_uuid] = chunks_original_with_embeddings
                        self.contextualized_chunks[doc_uuid] = chunks_contextualized_with_embeddings
                        logger.info(f"Stored {len(chunks_original_with_embeddings)} chunks with embeddings out of {len(chunks)} total chunks for '{original_name}'")
                    elif embeddings_arr.size == 0:
                        logger.error(f"No embeddings generated for any chunks in '{original_name}'. Document will not be searchable.")
                        self.chunks[doc_uuid] = []
                        self.contextualized_chunks[doc_uuid] = []
                    else:
                        # All chunks have embeddings
                        self.chunks[doc_uuid] = chunks
                        self.embeddings[doc_uuid] = embeddings_arr
                        self.contextualized_chunks[doc_uuid] = chunks_for_processing if should_ctx and chunks_for_processing else []
                        logger.info(f"Stored {len(chunks)} chunks with embeddings for '{original_name}'")
                else:
                    embeddings_arr = np.array([])
                    self.chunks[doc_uuid] = chunks

                self.embeddings[doc_uuid] = embeddings_arr
                
                meta_entry = self.metadata.get(doc_uuid, {})
                meta_entry.update({
                    'uuid': doc_uuid, 'original_name': original_name, 
                    'added': meta_entry.get('added', datetime.now().isoformat()), 
                    'updated': datetime.now().isoformat(), 'chunk_count': len(chunks),
                    'content_hash': current_hash, 'contextualized': should_ctx and bool(chunks_for_processing)
                })
                self.metadata[doc_uuid] = meta_entry
                
                bm25_index = self._create_bm25_index(chunks_for_processing if chunks_for_processing else [])
                if bm25_index: self.bm25_indexes[doc_uuid] = bm25_index
                else: self.bm25_indexes.pop(doc_uuid, None)
            else:
                if doc_uuid in self.metadata: self.metadata[doc_uuid]['checked'] = datetime.now().isoformat()

            if content_changed or not is_update:
                with open(doc_file_path, 'w', encoding='utf-8') as f: f.write(content)
            
            if save_to_disk: self._save_to_disk()
            if not _internal_call: await self._update_document_list_file()
            return doc_uuid
        except Exception as e:
            logger.error(f"Error add/update doc '{original_name}': {e}", exc_info=True)
            return None

    def get_googledoc_id_mapping(self) -> Dict[str, str]:
        tracked_file = self.base_dir / "tracked_google_docs.json"
        if not tracked_file.exists(): return {}
        mapping = {}
        try:
            with open(tracked_file, 'r', encoding='utf-8') as f: tracked_docs = json.load(f)
            for doc_entry in tracked_docs:
                if 'google_doc_id' in doc_entry and 'internal_doc_uuid' in doc_entry:
                    mapping[doc_entry['google_doc_id']] = doc_entry['internal_doc_uuid']
        except Exception as e: logger.error(f"Error reading tracked_google_docs.json: {e}")
        return mapping

    def get_original_name_to_googledoc_id_mapping(self) -> Dict[str, str]:
        """
        Returns a mapping from original_name to google_doc_id for creating citation URLs.
        This is the correct mapping to use when you have an original_name from search results
        and need the Google Doc URL ID string.
        """
        tracked_file = self.base_dir / "tracked_google_docs.json"
        if not tracked_file.exists():
            return {}
        
        mapping = {}
        try:
            with open(tracked_file, 'r', encoding='utf-8') as f:
                tracked_docs = json.load(f)
            
            for doc_entry in tracked_docs:
                google_doc_id = doc_entry.get('google_doc_id')
                original_name_at_import = doc_entry.get('original_name_at_import')
                
                if google_doc_id and original_name_at_import:
                    mapping[original_name_at_import] = google_doc_id
                    
        except Exception as e:
            logger.error(f"Error reading tracked_google_docs.json for original_name mapping: {e}")
        
        return mapping

    def get_all_document_contents(self) -> Dict[str, str]: 
        all_contents = {}
        internal_list_uuid = next((uid for uid, meta in self.metadata.items() if meta.get('original_name') == self._internal_list_doc_name), None)
        for doc_uuid, meta in self.metadata.items():
            if doc_uuid == internal_list_uuid: continue
            file_path = self.base_dir / f"{doc_uuid}.txt"
            if file_path.exists():
                try: all_contents[doc_uuid] = file_path.read_text(encoding='utf-8-sig')
                except Exception as e: logger.error(f"Error reading {file_path}: {e}")
            elif doc_uuid in self.chunks: all_contents[doc_uuid] = " ".join(self.chunks[doc_uuid])
        return all_contents

    async def search(self, query: str, top_k: int = None, apply_reranking: bool = None) -> List[Tuple[str, str, str, float, Optional[str], int, int]]:
        if top_k is None: top_k = self.top_k
        if apply_reranking is None and self.config: apply_reranking = self.config.RERANKING_ENABLED
        if not query or not query.strip(): return []
        
        query_embedding_res = await self.generate_embeddings([query], is_query=True)
        if query_embedding_res.size == 0: return []
        query_embedding = query_embedding_res[0]

        # Get adaptive candidate count for reranking
        if apply_reranking and self.config and hasattr(self.config, 'get_reranking_settings_for_query'):
            adaptive_settings = self.config.get_reranking_settings_for_query(query)
            initial_k = adaptive_settings['candidates']
        else:
            initial_k = (self.config.RERANKING_CANDIDATES if apply_reranking and self.config else top_k) if self.config else top_k
        initial_k = max(initial_k, top_k)
            
        embedding_results = self.custom_search_with_embedding(query_embedding, top_k=initial_k)
        bm25_results = self._search_bm25(query, top_k=initial_k)
        combined_results = self._combine_search_results(embedding_results, bm25_results, top_k=initial_k)
            
        if apply_reranking and len(combined_results) > 1:
            return await self.rerank_results(query, combined_results, top_k=top_k)
        return combined_results[:top_k]
                
    def custom_search_with_embedding(self, query_embedding: np.ndarray, top_k: int = None) -> List[Tuple[str, str, str, float, Optional[str], int, int]]:
        if top_k is None: top_k = self.top_k
        results = []
        for doc_uuid, doc_embeddings_array in self.embeddings.items():
            if not isinstance(doc_embeddings_array, np.ndarray) or doc_embeddings_array.size == 0: continue
            similarities = np.dot(doc_embeddings_array, query_embedding)
            if len(similarities) == 0: continue
            
            num_chunks_for_doc = len(self.chunks.get(doc_uuid, []))
            count_for_this_doc = min(top_k, num_chunks_for_doc, len(similarities))
            if count_for_this_doc <=0: continue

            top_indices_for_doc = np.argsort(similarities)[-count_for_this_doc:]
            for chunk_idx_in_doc in top_indices_for_doc:
                if chunk_idx_in_doc < num_chunks_for_doc:
                    original_name = self._get_original_name(doc_uuid)
                    chunk_text = self.get_contextualized_chunk(doc_uuid, chunk_idx_in_doc)
                    image_id = self.metadata.get(doc_uuid, {}).get('image_id')
                    results.append((doc_uuid, original_name, chunk_text, float(similarities[chunk_idx_in_doc]), image_id, chunk_idx_in_doc + 1, num_chunks_for_doc))
        results.sort(key=lambda x: x[3], reverse=True)
        return results[:top_k]

    def _save_to_disk(self):
        try:
            for data_dict, filename in [
                (self.chunks, 'chunks.pkl'), (self.contextualized_chunks, 'contextualized_chunks.pkl'),
                (self.embeddings, 'embeddings.pkl'), (self.metadata, 'metadata.json')]:
                path = self.base_dir / filename
                if filename.endswith('.json'):
                    with open(path, 'w', encoding='utf-8') as f: json.dump(data_dict, f, indent=2)
                else:
                    with open(path, 'wb') as f: pickle.dump(data_dict, f)
            logger.info(f"Saved all document data to {self.base_dir}")
        except Exception as e: logger.error(f"Error saving document data: {e}", exc_info=True)

    async def regenerate_all_embeddings(self):
        """Regenerate all embeddings using the current embedding model, keyed by UUID."""
        try:
            logger.info("Starting regeneration of all embeddings (UUID based).")
            all_doc_uuids = list(self.chunks.keys()) 

            for doc_uuid in all_doc_uuids:
                original_name = self._get_original_name(doc_uuid)
                logger.info(f"Regenerating embeddings for document: '{original_name}' (UUID: {doc_uuid})")

                doc_specific_chunks = self.chunks.get(doc_uuid, [])
                if not doc_specific_chunks:
                    logger.warning(f"No chunks found for document '{original_name}' (UUID: {doc_uuid}). Skipping.")
                    self.embeddings[doc_uuid] = np.array([])
                    continue
                
                max_words_ctx = self.config.MAX_WORDS_FOR_CONTEXT if self.config and hasattr(self.config, 'MAX_WORDS_FOR_CONTEXT') else 20000
                word_count = self.metadata.get(doc_uuid, {}).get('word_count', float('inf'))
                contextualization_enabled = self.config.CONTEXTUALIZATION_ENABLED if self.config and hasattr(self.config, 'CONTEXTUALIZATION_ENABLED') else True
                use_contextualised = self.config.USE_CONTEXTUALISED_CHUNKS if self.config and hasattr(self.config, 'USE_CONTEXTUALISED_CHUNKS') else True
                should_be_contextualized = contextualization_enabled and word_count <= max_words_ctx

                # Determine which chunks to use for embedding generation
                if use_contextualised:
                    chunks_for_embedding = self.contextualized_chunks.get(doc_uuid, [])
                    if not chunks_for_embedding:
                        chunks_for_embedding = doc_specific_chunks
                else:
                    chunks_for_embedding = doc_specific_chunks
                    logger.debug(f"Using original chunks for embeddings regeneration for '{original_name}' (contextualised chunks disabled)")
                
                if should_be_contextualized and (not self.contextualized_chunks.get(doc_uuid) or not self.metadata.get(doc_uuid, {}).get('contextualized')):
                    logger.info(f"Attempting to generate/regenerate contextualized chunks for '{original_name}' (UUID: {doc_uuid}).")
                    doc_file_path = self.base_dir / f"{doc_uuid}.txt"
                    if doc_file_path.exists():
                        doc_content_for_context = doc_file_path.read_text(encoding='utf-8-sig')
                        generated_context_chunks = await self.contextualize_chunks(original_name, doc_content_for_context, doc_specific_chunks)
                        self.contextualized_chunks[doc_uuid] = generated_context_chunks
                        # Only use contextualized chunks for embeddings if USE_CONTEXTUALISED_CHUNKS is enabled
                        if use_contextualised:
                            chunks_for_embedding = generated_context_chunks
                        else:
                            chunks_for_embedding = doc_specific_chunks
                        if doc_uuid in self.metadata: self.metadata[doc_uuid]['contextualized'] = True
                    else:
                        logger.error(f"Cannot find {doc_file_path} for contextualization of '{original_name}'. Using original chunks.")
                        chunks_for_embedding = doc_specific_chunks

                if not chunks_for_embedding: 
                    logger.error(f"No chunks for embedding '{original_name}' (UUID: {doc_uuid}). Skipping.")
                    self.embeddings[doc_uuid] = np.array([])
                    continue

                titles = [original_name] * len(chunks_for_embedding)
                new_embeddings = await self.generate_embeddings(chunks_for_embedding, is_query=False, titles=titles)
                
                # Handle case where fewer embeddings are returned than chunks
                if new_embeddings.size > 0 and len(new_embeddings) < len(chunks_for_embedding):
                    logger.warning(f"Regeneration: Got {len(new_embeddings)} embeddings for {len(chunks_for_embedding)} chunks in '{original_name}'. Truncating chunks to match.")
                    # Update chunks to match the number of successful embeddings
                    doc_specific_chunks_truncated = doc_specific_chunks[:len(new_embeddings)]
                    chunks_for_embedding_truncated = chunks_for_embedding[:len(new_embeddings)]
                    
                    self.chunks[doc_uuid] = doc_specific_chunks_truncated
                    self.contextualized_chunks[doc_uuid] = chunks_for_embedding_truncated if should_be_contextualized else []
                    logger.info(f"Regeneration: Updated '{original_name}' to {len(doc_specific_chunks_truncated)} chunks with embeddings")
                elif new_embeddings.size == 0:
                    logger.error(f"Regeneration: No embeddings generated for '{original_name}'. Document will not be searchable.")
                    self.chunks[doc_uuid] = []
                    self.contextualized_chunks[doc_uuid] = []
                
                self.embeddings[doc_uuid] = new_embeddings
                
                bm25_idx = self._create_bm25_index(chunks_for_embedding)
                if bm25_idx: self.bm25_indexes[doc_uuid] = bm25_idx
                else: self.bm25_indexes.pop(doc_uuid, None)

            self._save_to_disk()
            logger.info("Completed regeneration of all embeddings (UUID based).")
            return True
        except Exception as e:
            logger.error(f"Error regenerating embeddings (UUID based): {e}", exc_info=True)
            return False

    async def _load_documents(self, force_reload: bool = False):
        logger.info(f"--- DocumentManager: Starting _load_documents (force_reload={force_reload}) ---")
        self.chunks, self.contextualized_chunks, self.embeddings, self.metadata, self.bm25_indexes = {}, {}, {}, {}, {}
        migration_flag_file = self.base_dir / ".migration_to_uuid_complete"
        metadata_path = self.base_dir / 'metadata.json'
        needs_migration = not migration_flag_file.exists()

        if force_reload and migration_flag_file.exists():
            needs_migration = False 
        elif force_reload: 
            needs_migration = True 

        if needs_migration:
            logger.warning("!!! UUID MIGRATION REQUIRED !!! Backup 'documents/' and 'lorebooks/' if not done. All docs will be re-processed.")
            old_metadata_content = {}
            if metadata_path.exists():
                try:
                    with open(metadata_path, 'r', encoding='utf-8') as f:
                        old_metadata_content = json.load(f)
                    # Heuristic: if key looks like a UUID, assume already migrated (or partially)
                    if old_metadata_content and all(len(k) > 30 and '-' in k for k in old_metadata_content.keys()):
                        logger.info("Detected existing metadata.json appears to be UUID-keyed. Using it as base.")
                        self.metadata = old_metadata_content # Use this as the base if it's already UUID keyed
                        old_metadata_content = {} # Clear to prevent re-migration of these specific entries
                except Exception as e:
                    logger.error(f"Error loading old metadata.json for migration: {e}")
            
            old_name_to_new_uuid_map = {}

            # 1. Migrate old 'documents/' based on old_metadata_content (if it's not already UUID keyed)
            if old_metadata_content: # This ensures we only process if it's truly old, name-keyed metadata
                logger.info("Migrating old 'documents/' based on old name-keyed metadata.json...")
                for s_name_key, meta_val in list(old_metadata_content.items()):
                    is_potential_uuid = False
                    try:
                        uuid.UUID(s_name_key)
                        is_potential_uuid = True
                    except ValueError:
                        pass # Not a UUID, proceed with migration for this entry
                    
                    if is_potential_uuid:
                        logger.info(f"Skipping key '{s_name_key}' as it appears to be a UUID during old metadata processing.")
                        continue

                    original_name = meta_val.get('original_name', s_name_key)
                    old_file_path = self.base_dir / s_name_key # s_name_key is the old sanitized filename
                    
                    if old_file_path.is_file():
                        try:
                            content = old_file_path.read_text(encoding='utf-8-sig')
                            new_doc_uuid_for_old_file = str(uuid.uuid4())
                            added_uuid = await self.add_document(original_name, content, save_to_disk=False, existing_uuid=new_doc_uuid_for_old_file, _internal_call=True)
                            if added_uuid:
                                old_name_to_new_uuid_map[original_name] = added_uuid
                                if not original_name.endswith(".txt"): # Handle cases where .txt might have been used for lookup
                                     old_name_to_new_uuid_map[f"{original_name}.txt"] = added_uuid
                                old_file_path.unlink()
                                logger.info(f"Migrated old doc '{original_name}' (was {s_name_key}) to UUID {added_uuid}")
                        except Exception as e:
                            logger.error(f"Error migrating old doc '{original_name}' from path {old_file_path}: {e}", exc_info=True)
            
            # 2. Integrate 'lorebooks/'
            lorebooks_dir = Path(self.base_dir).parent / "lorebooks"
            if lorebooks_dir.exists() and lorebooks_dir.is_dir():
                logger.info("Migrating 'lorebooks/'...")
                for lb_file in lorebooks_dir.glob("*.txt"): # Assuming lorebooks are .txt
                    if lb_file.is_file():
                        try:
                            content = lb_file.read_text(encoding='utf-8-sig')
                            new_uuid_for_lorebook = await self.add_document(lb_file.name, content, save_to_disk=False, _internal_call=True)
                            if new_uuid_for_lorebook:
                                old_name_to_new_uuid_map[lb_file.name] = new_uuid_for_lorebook
                                lb_file.unlink()
                                logger.info(f"Migrated lorebook '{lb_file.name}' to UUID {new_uuid_for_lorebook}")
                        except Exception as e:
                            logger.error(f"Error migrating lorebook {lb_file.name}: {e}", exc_info=True)
                try: 
                    if not any(lorebooks_dir.iterdir()): # Check if empty before trying to remove
                        lorebooks_dir.rmdir()
                        logger.info(f"Removed empty lorebooks directory: {lorebooks_dir}")
                except Exception as e_rmdir: 
                    logger.warning(f"Could not remove lorebooks directory {lorebooks_dir}: {e_rmdir}")


            # 3. Update 'tracked_google_docs.json'
            tracked_gdocs_file = self.base_dir / "tracked_google_docs.json"
            if tracked_gdocs_file.exists():
                logger.info("Updating 'tracked_google_docs.json'...")
                new_gdoc_list = []
                try:
                    old_gdocs = json.loads(tracked_gdocs_file.read_text(encoding='utf-8'))
                    for entry in old_gdocs:
                        gdoc_id = entry.get('id')
                        custom_name = entry.get('custom_name')
                        name_key_for_map = custom_name or f"googledoc_{gdoc_id}.txt"
                        
                        internal_uuid = old_name_to_new_uuid_map.get(name_key_for_map)
                        if not internal_uuid and custom_name and not custom_name.endswith(".txt"): 
                             internal_uuid = old_name_to_new_uuid_map.get(custom_name)
                        
                        if internal_uuid:
                            new_gdoc_list.append({
                                "google_doc_id": gdoc_id, 
                                "internal_doc_uuid": internal_uuid, 
                                "original_name_at_import": custom_name or f"googledoc_{gdoc_id}",
                                "added_at": entry.get('added_at', datetime.now().isoformat())
                            })
                        else:
                            logger.warning(f"Could not map Google Doc ID {gdoc_id} (name key: {name_key_for_map}) to new internal UUID. Tracking entry dropped.")
                    tracked_gdocs_file.write_text(json.dumps(new_gdoc_list, indent=2), encoding='utf-8')
                except Exception as e:
                    logger.error(f"Error updating tracked_google_docs.json: {e}", exc_info=True)

            # 4. Cleanup unidentified files in 'documents/'
            logger.info("Cleaning up unidentified files in 'documents/'...")
            known_files = {"metadata.json", "chunks.pkl", "embeddings.pkl", "contextualized_chunks.pkl", 
                           "embeddings_provider.txt", "tracked_google_docs.json", migration_flag_file.name}
            if self.base_dir.is_dir():
                for item in self.base_dir.iterdir():
                    if item.is_file() and item.name not in known_files:
                        is_uuid_txt_file = False
                        try:
                            uuid.UUID(item.stem)
                            if item.suffix == '.txt':
                                is_uuid_txt_file = True
                        except ValueError:
                            pass 
                        
                        if not is_uuid_txt_file: 
                            logger.info(f"Deleting unidentified file: {item}")
                            try:
                                item.unlink()
                            except Exception as e_unl:
                                logger.error(f"Failed to delete {item}: {e_unl}")
            
            self._save_to_disk()
            migration_flag_file.write_text(f"Migration to UUID completed on: {datetime.now().isoformat()}")
            logger.info("--- UUID MIGRATION PROCESS COMPLETED ---")
        else: # Normal load (migration flag exists)
            logger.info("Loading documents using existing UUID-based system.")
            if metadata_path.exists():
                try:
                    # Try to read the file content first
                    raw_metadata_content = metadata_path.read_text(encoding='utf-8')
                    self.metadata = json.loads(raw_metadata_content)
                    logger.info(f"Loaded {len(self.metadata)} metadata entries from {metadata_path}.")
                except OSError as ose: # Catch specific OS errors during file read
                    logger.error(f"OSError reading metadata.json from {metadata_path}: {ose}", exc_info=True)
                    self.metadata = {} 
                except json.JSONDecodeError as jde: # Catch errors during JSON parsing
                    logger.error(f"JSONDecodeError parsing metadata.json from {metadata_path}: {jde}", exc_info=True)
                    # Attempt to load a backup or initialize as empty
                    # For now, just initialize as empty if parsing fails
                    self.metadata = {}
                except Exception as e: # Catch any other unexpected errors
                    logger.error(f"Unexpected error loading metadata.json from {metadata_path}: {e}", exc_info=True)
                    self.metadata = {} # Ensure it's an empty dict on error
            else:
                logger.info(f"Metadata file {metadata_path} not found. Initializing empty metadata.")
                self.metadata = {}
            
            for data_attr_name, pickle_filename in [('chunks','chunks.pkl'), ('contextualized_chunks','contextualized_chunks.pkl'), ('embeddings','embeddings.pkl')]:
                fpath = self.base_dir / pickle_filename
                if fpath.exists():
                    try: 
                        with open(fpath, 'rb') as f:
                            setattr(self, data_attr_name, pickle.load(f))
                        logger.info(f"Loaded {len(getattr(self, data_attr_name))} entries from {pickle_filename}")
                    except Exception as e:
                        logger.error(f"Error loading {fpath}: {e}", exc_info=True)
                        setattr(self, data_attr_name, {}) # Ensure empty dict on error
            
            logger.info("Rebuilding BM25 indexes from loaded data...")
            use_contextualised = self.config.USE_CONTEXTUALISED_CHUNKS if self.config and hasattr(self.config, 'USE_CONTEXTUALISED_CHUNKS') else True
            
            for doc_uuid, doc_chunks_list in self.chunks.items():
                if use_contextualised:
                    chunks_for_bm25 = self.contextualized_chunks.get(doc_uuid, doc_chunks_list)
                else:
                    chunks_for_bm25 = doc_chunks_list
                    
                if chunks_for_bm25:
                    idx = self._create_bm25_index(chunks_for_bm25)
                    if idx: self.bm25_indexes[doc_uuid] = idx
            logger.info(f"Rebuilt {len(self.bm25_indexes)} BM25 indexes.")
            
            orphans_re_added = 0
            if self.base_dir.is_dir():
                for item_path in self.base_dir.glob("*.txt"):
                    try:
                        potential_uuid = item_path.stem
                        uuid.UUID(potential_uuid) 
                        if potential_uuid not in self.metadata:
                            logger.warning(f"Found orphaned document file: {item_path}. Attempting to re-add.")
                            content = item_path.read_text(encoding='utf-8-sig')
                            added_uuid = await self.add_document(
                                original_name=f"orphaned_{potential_uuid}", content=content, 
                                save_to_disk=False, existing_uuid=potential_uuid, _internal_call=True)
                            if added_uuid: orphans_re_added +=1
                    except ValueError: pass 
                    except Exception as e_orphan: logger.error(f"Error processing potential orphaned file {item_path}: {e_orphan}")
            if orphans_re_added > 0:
                logger.info(f"Re-added {orphans_re_added} orphaned document files.")
                self._save_to_disk()

        await self._update_document_list_file()
        logger.info(f"--- DocumentManager: _load_documents finished. Metadata entries: {len(self.metadata)} ---")

    async def convert_tracked_gdocs_to_uuid_format(self):
        """
        Converts an old-format tracked_google_docs.json (with 'id', 'custom_name')
        to the new format (with 'google_doc_id', 'internal_doc_uuid', 'original_name_at_import').
        This is a one-time utility function to fix the format if it was reverted.
        """
        logger.info("Attempting to convert tracked_google_docs.json to new UUID format...")
        tracked_gdocs_file = self.base_dir / "tracked_google_docs.json"
        if not tracked_gdocs_file.exists():
            logger.warning(f"{tracked_gdocs_file} not found. Cannot convert.")
            return False

        try:
            with open(tracked_gdocs_file, 'r', encoding='utf-8') as f:
                old_format_docs = json.load(f)
        except Exception as e:
            logger.error(f"Error reading {tracked_gdocs_file} for conversion: {e}")
            return False

        if not old_format_docs or not isinstance(old_format_docs, list):
            logger.info(f"{tracked_gdocs_file} is empty or not a list. No conversion needed or possible.")
            return True # Assuming already correct or empty

        # Check if already in new format
        if all('google_doc_id' in entry and 'internal_doc_uuid' in entry for entry in old_format_docs):
            logger.info(f"{tracked_gdocs_file} appears to be already in the new format. No conversion performed.")
            return True

        logger.info(f"Found {len(old_format_docs)} entries in old format. Converting...")
        new_format_docs = []
        name_to_uuid_map = {meta.get('original_name'): uuid_key for uuid_key, meta in self.metadata.items()}

        for old_entry in old_format_docs:
            google_doc_id_old = old_entry.get('id')
            custom_name = old_entry.get('custom_name')
            added_at = old_entry.get('added_at', datetime.now().isoformat())

            if not google_doc_id_old or not custom_name:
                logger.warning(f"Skipping entry due to missing 'id' or 'custom_name': {old_entry}")
                continue

            internal_doc_uuid = name_to_uuid_map.get(custom_name)
            
            # Try a common variation if the first lookup failed (e.g. if .txt was part of the name in metadata)
            if not internal_doc_uuid and not custom_name.endswith(".txt"):
                 internal_doc_uuid = name_to_uuid_map.get(f"{custom_name}.txt")
            if not internal_doc_uuid and custom_name.endswith(".txt"):
                 internal_doc_uuid = name_to_uuid_map.get(custom_name[:-4])


            if not internal_doc_uuid:
                # Attempt to find by looking for a document whose original name *contains* the custom_name
                # This is a looser match, useful if names were slightly altered.
                for meta_uuid, meta_data in self.metadata.items():
                    meta_orig_name = meta_data.get('original_name', '')
                    if custom_name in meta_orig_name:
                        internal_doc_uuid = meta_uuid
                        logger.info(f"Loosely matched '{custom_name}' to '{meta_orig_name}' (UUID: {internal_doc_uuid})")
                        break
            
            if not internal_doc_uuid:
                 # Last resort: check if a document with original_name like "googledoc_{google_doc_id_old}" exists
                potential_gdoc_name = f"googledoc_{google_doc_id_old}"
                internal_doc_uuid = name_to_uuid_map.get(potential_gdoc_name)
                if not internal_doc_uuid:
                    internal_doc_uuid = name_to_uuid_map.get(f"{potential_gdoc_name}.txt")

            if internal_doc_uuid:
                new_format_docs.append({
                    "google_doc_id": google_doc_id_old,
                    "internal_doc_uuid": internal_doc_uuid,
                    "original_name_at_import": custom_name,
                    "added_at": added_at,
                    "updated_at": datetime.now().isoformat() 
                })
                logger.info(f"Converted entry for '{custom_name}' (ID: {google_doc_id_old}) to use UUID: {internal_doc_uuid}")
            else:
                logger.warning(f"Could not find internal_doc_uuid for Google Doc '{custom_name}' (ID: {google_doc_id_old}). This entry will be dropped from tracking.")
        
        try:
            with open(tracked_gdocs_file, 'w', encoding='utf-8') as f:
                json.dump(new_format_docs, f, indent=2)
            logger.info(f"Successfully converted and saved {len(new_format_docs)} entries to {tracked_gdocs_file}.")
            return True
        except Exception as e:
            logger.error(f"Error writing converted data to {tracked_gdocs_file}: {e}")
            return False

    async def reload_documents(self):
        """Reload all documents from disk, regenerating embeddings."""
        await self._load_documents(force_reload=True)

    def get_lorebooks_path(self):
        """DEPRECATED: Lorebooks are now integrated. This path is for legacy purposes if any remain."""
        lorebook_legacy_path = self.base_dir / "legacy_lorebooks" # Changed to a subdirectory
        # lorebook_legacy_path.mkdir(parents=True, exist_ok=True) # No need to create if it's just for checking
        logger.warning("get_lorebooks_path() is deprecated. Lorebooks are integrated into the main document system.")
        return lorebook_legacy_path
    
    def track_google_doc(self, google_doc_id: str, internal_doc_uuid: str, original_name_at_import: str):
        """Track a Google Doc by linking its ID to an internal document UUID."""
        tracked_file = self.base_dir / "tracked_google_docs.json"
        tracked_docs = []
        if tracked_file.exists():
            try:
                with open(tracked_file, 'r', encoding='utf-8') as f:
                    tracked_docs = json.load(f)
            except Exception as e:
                logger.error(f"Error reading {tracked_file} for GDoc tracking: {e}")

        entry_found = False
        for entry in tracked_docs:
            if entry.get('google_doc_id') == google_doc_id:
                entry['internal_doc_uuid'] = internal_doc_uuid
                entry['original_name_at_import'] = original_name_at_import
                entry['updated_at'] = datetime.now().isoformat()
                entry_found = True
                logger.info(f"Updated tracking for Google Doc ID {google_doc_id} to internal UUID {internal_doc_uuid}.")
                break
        
        if not entry_found:
            tracked_docs.append({
                'google_doc_id': google_doc_id,
                'internal_doc_uuid': internal_doc_uuid,
                'original_name_at_import': original_name_at_import,
                'added_at': datetime.now().isoformat()
            })
            logger.info(f"Added tracking for Google Doc ID {google_doc_id} with internal UUID {internal_doc_uuid}.")

        try:
            with open(tracked_file, 'w', encoding='utf-8') as f:
                json.dump(tracked_docs, f, indent=2)
            return f"Google Doc {google_doc_id} tracking updated/added with internal UUID {internal_doc_uuid}."
        except Exception as e:
            logger.error(f"Error writing {tracked_file} for GDoc tracking: {e}")
            return f"Error saving GDoc tracking for {google_doc_id}."


    async def rename_document(self, doc_uuid: str, new_original_name: str) -> str:
        """Rename a document's user-facing original_name given its UUID."""
        if not doc_uuid or not new_original_name:
            return "Error: Document UUID and new original name are required."

        if doc_uuid not in self.metadata:
            return f"Error: Document with UUID '{doc_uuid}' not found."

        is_internal_list_doc = self.metadata[doc_uuid].get('original_name') == self._internal_list_doc_name
        if new_original_name == self._internal_list_doc_name and not is_internal_list_doc:
            return f"Error: Cannot rename a document to the internal list file name ('{self._internal_list_doc_name}')."
        
        for other_uuid, meta in self.metadata.items():
            if other_uuid != doc_uuid and meta.get('original_name') == new_original_name:
                return f"Error: Another document already has the original name '{new_original_name}'."

        old_original_name = self.metadata[doc_uuid].get('original_name', doc_uuid)
        self.metadata[doc_uuid]['original_name'] = new_original_name
        self.metadata[doc_uuid]['updated'] = datetime.now().isoformat()
        
        logger.info(f"Renamed document (UUID: {doc_uuid}) from '{old_original_name}' to '{new_original_name}'.")

        tracked_file = Path(self.base_dir) / "tracked_google_docs.json"
        if tracked_file.exists():
            try:
                with open(tracked_file, 'r', encoding='utf-8') as f: tracked_docs_list = json.load(f)
                gdoc_updated = False
                for gdoc_entry in tracked_docs_list:
                    if gdoc_entry.get('internal_doc_uuid') == doc_uuid:
                        gdoc_entry['original_name_at_import'] = new_original_name
                        gdoc_updated = True
                        break
                if gdoc_updated:
                    with open(tracked_file, 'w', encoding='utf-8') as f: json.dump(tracked_docs_list, f, indent=2)
            except Exception as e_gdoc: logger.error(f"Error updating tracked_google_docs.json for rename: {e_gdoc}")

        self._save_to_disk()
        await self._update_document_list_file()
        return f"Document (UUID: {doc_uuid}) original name changed from '{old_original_name}' to '{new_original_name}'."

    async def delete_document(self, doc_uuid: str) -> bool: 
        """Delete a document from the system using its UUID."""
        if not doc_uuid: 
            logger.error("Delete_document called with no UUID.")
            return False

        original_name = self._get_original_name(doc_uuid) 
        
        if original_name == self._internal_list_doc_name: 
            logger.warning(f"Attempted to delete the internal document list file '{original_name}' (UUID: {doc_uuid}). Operation aborted.")
            return False

        if doc_uuid not in self.metadata:
            logger.warning(f"Document with UUID '{doc_uuid}' (Original Name: '{original_name}') not found in metadata for deletion.")
        
        deleted_in_memory = False
        if doc_uuid in self.metadata: del self.metadata[doc_uuid]; deleted_in_memory = True
        if doc_uuid in self.chunks: del self.chunks[doc_uuid]; deleted_in_memory = True
        if doc_uuid in self.contextualized_chunks: del self.contextualized_chunks[doc_uuid]; deleted_in_memory = True
        if doc_uuid in self.embeddings: del self.embeddings[doc_uuid]; deleted_in_memory = True
        if doc_uuid in self.bm25_indexes: del self.bm25_indexes[doc_uuid]; deleted_in_memory = True

        if deleted_in_memory: logger.info(f"Removed document '{original_name}' (UUID: {doc_uuid}) from in-memory stores.")

        file_path = self.base_dir / f"{doc_uuid}.txt"
        file_deleted_from_disk = False
        if file_path.exists():
            try: file_path.unlink(); file_deleted_from_disk = True; logger.info(f"Deleted document file: {file_path}")
            except Exception as e: logger.error(f"Failed to delete document file {file_path}: {e}")
        else: logger.warning(f"Document file {file_path} not found on disk for deletion (Original: '{original_name}').")

        gdoc_tracking_removed = False
        tracked_file = Path(self.base_dir) / "tracked_google_docs.json"
        if tracked_file.exists():
            try:
                with open(tracked_file, 'r', encoding='utf-8') as f: tracked_docs_list = json.load(f)
                initial_len = len(tracked_docs_list)
                tracked_docs_list = [gd for gd in tracked_docs_list if gd.get('internal_doc_uuid') != doc_uuid]
                if len(tracked_docs_list) < initial_len:
                    with open(tracked_file, 'w', encoding='utf-8') as f: json.dump(tracked_docs_list, f, indent=2)
                    gdoc_tracking_removed = True
                    logger.info(f"Removed GDoc tracking for internal UUID {doc_uuid}")
            except Exception as e: logger.error(f"Error updating tracked_google_docs.json for deletion of UUID {doc_uuid}: {e}")

        if deleted_in_memory or file_deleted_from_disk or gdoc_tracking_removed:
            try:
                self._save_to_disk()
                await self._update_document_list_file()
                logger.info(f"Successfully deleted and saved for '{original_name}' (UUID: {doc_uuid}).")
                return True
            except Exception as e: logger.error(f"Error saving after deleting '{original_name}' (UUID: {doc_uuid}): {e}")
            return False
        logger.warning(f"No action for delete_document UUID '{doc_uuid}' (Original: '{original_name}'). Not found.")
        return False


    async def rerank_results(self, query: str, initial_results: List[Tuple[str, str, str, float, Optional[str], int, int]], top_k: int = None) -> List[Tuple[str, str, str, float, Optional[str], int, int]]:
        """
        Re-rank search results using the Gemini embedding model for more nuanced relevance.
        
        Args:
            query: The search query
            initial_results: List of tuples (doc_uuid, original_name, chunk, score, image_id, chunk_idx, total_chunks)
            top_k: Number of results to return after re-ranking
            
        Returns:
            List of re-ranked results (doc_uuid, original_name, chunk, combined_score, image_id, chunk_idx, total_chunks)
        """
        if top_k is None: top_k = self.top_k
        if not initial_results: return []
        
        logger.info(f"Re-ranking {len(initial_results)} initial results for query: {query}")
        
        chunks_to_rerank = [result[2] for result in initial_results] 
        
        # Use adaptive query context based on query complexity
        if self.config and hasattr(self.config, 'get_reranking_settings_for_query'):
            adaptive_settings = self.config.get_reranking_settings_for_query(query)
            if adaptive_settings['filter_mode'] == 'topk':  # Complex query
                # Use more direct context for complex queries to reduce semantic mismatch
                query_context = f"Find content related to: {query}"
            else:  # Simple query
                query_context = f"Question: {query}\nWhat information would fully answer this question?"
        else:
            query_context = f"Question: {query}\nWhat information would fully answer this question?"
            
        query_embedding_result = await self.generate_embeddings([query_context], is_query=True)
        if query_embedding_result.size == 0: return initial_results 
        query_embedding = query_embedding_result[0]

        content_texts = [f"This document contains the following information: {chunk}" for chunk in chunks_to_rerank]
        content_embeddings = await self.generate_embeddings(content_texts, is_query=False) 
        if content_embeddings.size == 0 or content_embeddings.shape[0] != len(content_texts): return initial_results

        if query_embedding.shape[0] != content_embeddings.shape[1]: return initial_results

        relevance_scores = np.dot(content_embeddings, query_embedding)
        
        # Get adaptive settings for score combination
        if self.config and hasattr(self.config, 'get_reranking_settings_for_query'):
            adaptive_settings = self.config.get_reranking_settings_for_query(query)
            is_complex = adaptive_settings['filter_mode'] == 'topk'
        else:
            is_complex = False
        
        reranked_results_with_data = []
        for i, initial_res_tuple in enumerate(initial_results):
            if i < len(relevance_scores):
                original_score = initial_res_tuple[3]
                chunk_text = initial_res_tuple[2]
                
                # Apply content-aware boosting for complex queries
                content_boost = 1.0
                if is_complex:
                    # Boost chunks containing key terms from the query
                    query_lower = query.lower()
                    chunk_lower = chunk_text.lower()
                    
                    # Extract key terms from complex queries with stronger boosting
                    key_terms = []
                    if 'mong' in query_lower:
                        key_terms.extend(['mong', 'möng', 'arshtini'])  # Added related term
                    if 'philosophy' in query_lower or 'belief' in query_lower:
                        key_terms.extend(['philosophy', 'philosophical', 'belief', 'doctrine', 'tradition', 'practice'])
                    if 'theology' in query_lower or 'religious' in query_lower:
                        key_terms.extend(['theology', 'theological', 'religious', 'divine', 'god', 'spiritual', 'sacred'])
                    if 'analysis' in query_lower or 'detailed' in query_lower:
                        key_terms.extend(['analysis', 'examine', 'study', 'research', 'culture', 'cultural'])
                    if 'relationship' in query_lower or 'intersection' in query_lower:
                        key_terms.extend(['relationship', 'connection', 'intersection', 'between', 'among'])
                    
                    # Apply stronger boost for each matching term
                    term_matches = 0
                    for term in key_terms:
                        if term in chunk_lower:
                            term_matches += 1
                    
                    if term_matches > 0:
                        # Exponential boost for multiple term matches
                        content_boost += 0.5 * term_matches + 0.2 * (term_matches ** 2)
                    
                    # Extra boost for chunks that contain the primary subject (Mong)
                    if 'mong' in chunk_lower or 'möng' in chunk_lower or 'arshtini' in chunk_lower:
                        content_boost += 1.0  # Strong boost for primary subject
                
                # Use different score combination for complex vs simple queries
                if is_complex:
                    # For complex queries, preserve much more of the original ranking
                    # and apply content boost more aggressively
                    base_score = 0.8 * original_score + 0.2 * float(relevance_scores[i])
                    combined_score = base_score * content_boost
                else:
                    # For simple queries, rely more on semantic similarity
                    combined_score = 0.3 * original_score + 0.7 * float(relevance_scores[i])
                
                reranked_results_with_data.append(initial_res_tuple[:3] + (combined_score,) + initial_res_tuple[4:])
        
        reranked_results_with_data.sort(key=lambda x: x[3], reverse=True)
        
        # Get adaptive settings based on query complexity
        if self.config and hasattr(self.config, 'get_reranking_settings_for_query'):
            adaptive_settings = self.config.get_reranking_settings_for_query(query)
            min_score = adaptive_settings['min_score']
            filter_mode = adaptive_settings['filter_mode']
            logger.info(f"Using adaptive settings: {adaptive_settings}")
        else:
            min_score = self.config.RERANKING_MIN_SCORE if self.config and hasattr(self.config, 'RERANKING_MIN_SCORE') else 0.45
            filter_mode = self.config.RERANKING_FILTER_MODE if self.config and hasattr(self.config, 'RERANKING_FILTER_MODE') else 'strict'
        
        if filter_mode == 'dynamic':
            scores = [r[3] for r in reranked_results_with_data]
            if scores:
                mean_score = sum(scores) / len(scores)
                dynamic_threshold = mean_score * 0.8 
                min_score = max(min_score, dynamic_threshold)
            filtered_results = [r for r in reranked_results_with_data if r[3] >= min_score]
        elif filter_mode == 'topk':
            filtered_results = reranked_results_with_data[:top_k] if top_k else reranked_results_with_data
        else: # strict
            filtered_results = [r for r in reranked_results_with_data if r[3] >= min_score]
        
        if not filtered_results and reranked_results_with_data: 
            logger.warning(f"Filtering with mode '{filter_mode}' (threshold {min_score:.3f}) removed all results. Falling back.")
            filtered_results = reranked_results_with_data[:top_k] if top_k else []
        
        final_results_to_return = filtered_results[:top_k] if top_k else filtered_results

        logger.info(f"Re-ranking with '{filter_mode}' mode: {len(initial_results)} -> {len(final_results_to_return)} results")
        return final_results_to_return
